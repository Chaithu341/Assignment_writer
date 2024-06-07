import os
import google.generativeai as genai
import docx
import re

os.environ["GOOGLE_API_KEY"] = "AIzaSyDjqhJAOBeA0RmfSYnqDpCIn4ge-dclqRM"
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])

model = genai.GenerativeModel("gemini-pro")

def answer_question(question):
    response = model.generate_content(question)
    return response.text

def create_docx_file(question, answer, name, roll_no):
    desktop_dir = os.path.join(os.path.expanduser("~"), "Desktop")
    file_name = "q_and_a.docx"
    file_path = os.path.join(desktop_dir, file_name)

    doc = docx.Document()

    # Add border to the document
    section = doc.sections[0]
    section.top_margin = docx.shared.Inches(0.5)
    section.bottom_margin = docx.shared.Inches(0.5)
    section.left_margin = docx.shared.Inches(0.5)
    section.right_margin = docx.shared.Inches(0.5)

    # Add name and roll number to the top left side of the first page
    para = doc.add_paragraph("")
    run = para.add_run(f"{name}\n")
    font = run.font
    font.size = docx.shared.Pt(16)
    run = para.add_run(f"{roll_no}")
    font = run.font
    font.size = docx.shared.Pt(16)

    # Add question and answer
    doc.add_paragraph(f"Question: {question}")
    answer_paragraphs = answer.split("\n")
    for paragraph in answer_paragraphs:
        if re.match(r'\*{2}.*\*{2}', paragraph):
            para = doc.add_paragraph("")
            run = para.add_run(paragraph[2:-2])
            run.font.bold = True
        else:
            para = doc.add_paragraph(paragraph)

    doc.save(file_path)

    print(f"File created: {file_path}")

def main():
    name = "hi" #input("Enter your name: ")
    roll_no = "21" #input("Enter your roll number: ")
    question = "What is ai and ml? "  #input("Enter your question: ")
    answer = answer_question(question)
    print(answer)

    create_docx_file(question, answer, name, roll_no)

if __name__ == "__main__":
    main()